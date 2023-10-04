import csv
from docx import Document
import pandas as pd


def collect_user_data():
    '''This function collects user data as input and returns a dictionary of the userdata
        And writes distinct user data into a csv file'''
    try:
        # Accept user input
        name = input("Enter your name: ")
        age = int(input("Enter your age: "))
        gender = input("Enter your gender (Male/Female): ").capitalize()
        weight = int(input("Enter your weight (kg): "))
        height = float(input("Enter your height (cm): "))

        # Create a dictionary to store user data
        user_data = {
            "Name": name,
            "Age": age,
            "Gender": gender,
            "Weight (kg)": weight,
            "Height (cm)": height,
        }

        # Append user data to the CSV file
        with open("user_data.csv", mode="a", newline="") as file:
            writer = csv.DictWriter(file, fieldnames=user_data.keys())
            # Check if the file is empty, and if so, write the header row
            if file.tell() == 0:
                writer.writeheader()
            writer.writerow(user_data)

        return user_data
    except (ValueError, KeyboardInterrupt) as e:
        print(f"Error: {e}")
        return None


def calculate_metrics(user_data):
    '''This function calculates BMI, BMR, body fat percentage, and ideal weight range
    and returns a dictionary of the calculated metrics'''
    try:
        height_in_meters = user_data["Height (cm)"] / 100

        # Calculate BMI
        bmi = round(user_data["Weight (kg)"] / (height_in_meters ** 2), 2)

        # Calculate BMR
        if user_data["Gender"] == "Male":
            bmr = round(88.362 + (13.397 * user_data["Weight (kg)"]) + (4.799 * user_data["Height (cm)"]) - (
                    5.677 * user_data["Age"]), 2)
        elif user_data["Gender"] == "Female":
            bmr = round(447.593 + (9.247 * user_data["Weight (kg)"]) + (3.098 * user_data["Height (cm)"]) - (
                    4.330 * user_data["Age"]), 2)
        else:
            bmr = None  # Gender not recognized

        # Calculate body fat percentage and ideal weight range
        if user_data["Gender"] == "Male":
            body_fat_percentage = (1.20 * bmi) + (0.23 * user_data["Age"]) - 16.2
        elif user_data["Gender"] == "Female":
            body_fat_percentage = (1.20 * bmi) + (0.23 * user_data["Age"]) - 5.4
        else:
            body_fat_percentage = None

        if user_data["Gender"] == "Male":
            ideal_weight_low = round((user_data["Height (cm)"] - 100) - ((user_data["Height (cm)"] - 100) / 10),
                                     2)
            ideal_weight_high = round((user_data["Height (cm)"] - 100) + ((user_data["Height (cm)"] - 100) / 10),
                                      2)
        elif user_data["Gender"] == "Female":
            ideal_weight_low = round((user_data["Height (cm)"] - 100) + ((user_data["Height (cm)"] - 100) / 10),
                                     2)
            ideal_weight_high = round((user_data["Height (cm)"] - 100) - ((user_data["Height (cm)"] - 100) / 10),
                                      2)
        else:
            ideal_weight_low = None
            ideal_weight_high = None

        ideal_weight_range = ideal_weight_low, ideal_weight_high

        return {
            "BMR": bmr,
            "Body Fat Percentage": body_fat_percentage,
            "Ideal Weight Range (kg)": ideal_weight_range,
            "BMI": bmi,
        }
    except ZeroDivisionError as e:
        print(f"Error: {e}")
        return None


def categorize_user(user_data, calculate_metrics):
    '''This function categorizes the user based on the calculated metrics
    and returns a dictionary of the categories'''
    try:
        name = user_data["Name"]
        bmi = calculate_metrics["BMI"]
        bmr = calculate_metrics["BMR"]
        ideal_weight_range = calculate_metrics["Ideal Weight Range (kg)"]
        body_fat_percentage = calculate_metrics["Body Fat Percentage"]

        # BMI categories
        if bmi < 18.5:
            bmi_category = "Underweight"
        elif 18.5 <= bmi < 24.9:
            bmi_category = "Healthy"
        elif 24.9 <= bmi < 29.9:
            bmi_category = "Overweight"
        else:
            bmi_category = "Obese"

        # BMR categories
        if 1200 <= bmr < 1400:
            bmr_category = "Low BMR"
        elif 1400 <= bmr < 1800:
            bmr_category = "Moderate BMR"
        elif 1800 <= bmr < 2200:
            bmr_category = "High BMR"
        else:
            bmr_category = "Unknown"

        # Ideal weight range categories
        ideal_weight_low, ideal_weight_high = ideal_weight_range
        if user_data["Weight (kg)"] < ideal_weight_low:
            ideal_weight_category = "Underweight"
        elif user_data["Weight (kg)"] > ideal_weight_high:
            ideal_weight_category = "Overweight/Obese"
        else:
            ideal_weight_category = "Healthy"

        # Body fat percentage categories
        gender = user_data["Gender"]
        if gender == "Male":
            if body_fat_percentage < 6:
                body_fat_category = "Athletic/Lean"
            elif 6 <= body_fat_percentage <= 10:
                body_fat_category = "Healthy"
            else:
                body_fat_category = "Overweight/Obese"

        elif gender == "Female":
            if body_fat_percentage < 16:
                body_fat_category = "Athletic/Lean"
            elif 16 <= body_fat_percentage <= 20:
                body_fat_category = "Healthy Range"
            else:
                body_fat_category = "Overweight/Obese"
        else:
            body_fat_category = "Unknown"

        print(f"\n\nHi {name}!, Below is your weight category based on the details you provided. "
              f"\nYour Current weight category: '{bmi_category}'\nYour Ideal weight range: '{ideal_weight_range}'")

        return {
            "BMI Category": bmi_category,
            "BMR Category": bmr_category,
            "Ideal Weight (kg) Category": ideal_weight_category,
            "Body Fat Category": body_fat_category,
        }
    except (KeyError, TypeError) as e:
        print(f"Error: {e}")
        return None


def get_data_from_excel_category_wise(user_categories):
    '''This function gets diet and exercise data from Excel file based on the user category'''
    try:
        df = pd.read_excel('HealthMateAppData.xlsx')

        # Initialize empty dictionaries to store the results for each category
        exercise_results = {}
        diet_results = {}

        for metric, condition in user_categories.items():
            # Filter rows where "Metric" matches the metric and "Condition" matches the condition
            filtered_df = df[(df['Metric'] == metric) & (df['Condition'] == condition)]

            exercise_data = filtered_df['Exercise']

            diet_data = filtered_df['Diet']

            # Check if exercise_data and diet_data are not empty before adding them to the respective dictionaries
            if not exercise_data.empty:
                exercise_results[f'{metric} - {condition}'] = exercise_data.tolist()

            if not diet_data.empty:
                diet_results[f'{metric} - {condition}'] = diet_data.tolist()

        # Return the dictionaries containing exercise and diet data for each category
        return exercise_results, diet_results
    except (FileNotFoundError, KeyError) as e:
        print(f"Error: {e}")
        return None, None


def collect_goal():
    '''This function collects user goal and return the goal'''
    goals = []

    try:
        # Prompt the user to set goals
        set_goals = input("\nDo you want to set goals? (yes/no): ").strip().lower()

        if set_goals == "yes":
            print("Select a goal from the following options:")
            print("1. Lose weight/Fat")
            print("2. Gain weight/Muscle")
            print("3. Improve cardiovascular health")
            print("4. Maintain current weight")

            while True:
                try:
                    selected_goal = int(input("Enter the 'number' (1,2,3,4) of the goal you want to set: "))
                    if 1 <= selected_goal <= 4:
                        break  # Exit the loop if a valid goal is selected
                    else:
                        print("Invalid input. Please select a valid goal number.")
                except ValueError:
                    print("Invalid input. Please enter a number.")

            if selected_goal == 1:
                goals.append("Lose weight")
            elif selected_goal == 2:
                goals.append("Gain muscle")
            elif selected_goal == 3:
                goals.append("Improve cardiovascular health")
            elif selected_goal == 4:
                goals.append("Maintain current weight")
    except KeyboardInterrupt:
        print("\nGoal setting interrupted by the user.")

    return goals


def write_to_word_file(user_data, calculated_metrics, goals, exercise_data, diet_data):
    try:
        # Create a new Word document
        doc = Document()

        # Add a title to the document
        doc.add_heading('Health Report', 0)

        # Add User Data section
        doc.add_heading('User Data', level=1)
        for key, value in user_data.items():
            doc.add_paragraph(f'{key}: {value}')

        # Add Calculated Metrics section
        doc.add_heading('Calculated Metrics', level=1)
        for key, value in calculated_metrics.items():
            doc.add_paragraph(f'{key}: {value}')

        # Add Goals section
        doc.add_heading('Goals', level=1)
        for goal in goals:
            doc.add_paragraph(goal)

        # Add Exercise Data section
        doc.add_heading('Exercise Suggestions', level=1)
        for category, exercise_list in exercise_data.items():
            doc.add_heading(category, level=2)
            for exercise in exercise_list:
                doc.add_paragraph(exercise)

        # Add Diet Data section
        doc.add_heading('Diet Suggestions', level=1)
        for category, diet_list in diet_data.items():
            doc.add_heading(category, level=2)
            for diet in diet_list:
                doc.add_paragraph(diet)

        disclaimer = """
        All the health data is calculated using formulae available online. Health and Dietary suggestions are based on users category.
         Please consult one of our personal instructors before following the workout and diet plans.
        """
        doc.add_heading('*DISCLAIMER*', level=1)
        doc.add_paragraph(disclaimer)

        doc.save(f'{user_data["Name"]}_health_report.docx')
        print("\nCongrats! Your personalized Health report is now available. It includes your personalized Health and "
              "Dietary suggestions!!"
              "\nThe file name is saved with your name_health_report.docx")

    except Exception as e:
        print(f"An error occurred while saving the health report: {str(e)}")


if __name__ == "__main__":
    while True:
        user_data = collect_user_data()
        metrics = calculate_metrics(user_data)

        print("\nUser Data:")
        for key, value in user_data.items():
            print(f"{key}: {value}")
        print("\nBased on your inputs, here are some numbers which wont make sense to you: ")
        for key, value in metrics.items():
            print(f"{key}: {value}")

        category = categorize_user(user_data, metrics)
        goal = collect_goal()

        print("\nBased on each metric, below are your individual categories:")
        for key, value in category.items():
            print(f"{key}: {value}")

        if goal:
            print(f"\nYour goal is: {goal}")

        exercise_data, diet_data = get_data_from_excel_category_wise(category)
        write_to_word_file(user_data, metrics, goal, exercise_data, diet_data)

        # Ask the user if they want to enter new details
        another_entry = input("\nDo you want to enter new details? (yes/no): ").strip().lower()
        if another_entry != "yes":
            print("Thankyou for using HealthMate!!")
            break
